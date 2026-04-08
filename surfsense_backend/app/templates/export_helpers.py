"""
Helpers for report export templates.

* ``get_typst_template_path()``  - path to the custom Pandoc -> Typst template.
* ``get_html_css_path()``        - path to the CSS stylesheet for HTML exports.
* ``get_reference_docx_path()``  - path to a styled reference.docx for Pandoc.

The reference DOCX is generated lazily on first call from Pandoc's built-in
default, then restyled with *python-docx* and cached on disk so subsequent
exports are instant.
"""

from __future__ import annotations

import subprocess
import threading
from pathlib import Path

_DIR = Path(__file__).resolve().parent
_GENERATED_DIR = _DIR / "_generated"
_REFERENCE_DOCX = _GENERATED_DIR / "reference.docx"
_TYPST_TEMPLATE = _DIR / "report_pdf.typst"
_HTML_CSS = _DIR / "report_html.css"

_docx_lock = threading.Lock()


def get_typst_template_path() -> Path:
    return _TYPST_TEMPLATE


def get_html_css_path() -> Path:
    return _HTML_CSS


def get_reference_docx_path() -> Path:
    """Return path to the styled reference.docx, creating it if absent."""
    if _REFERENCE_DOCX.exists():
        return _REFERENCE_DOCX
    with _docx_lock:
        if _REFERENCE_DOCX.exists():
            return _REFERENCE_DOCX
        _generate_reference_docx()
    return _REFERENCE_DOCX


# ---------------------------------------------------------------------------
# Reference DOCX generation
# ---------------------------------------------------------------------------

_HEADING_COLOR_RGB = (0x1E, 0x29, 0x3B)  # Slate-900
_ACCENT_RGB = (0x25, 0x63, 0xEB)  # Blue-600


def _generate_reference_docx() -> None:
    """Build a professional reference.docx from Pandoc's default + restyling."""
    import pypandoc
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    _GENERATED_DIR.mkdir(parents=True, exist_ok=True)

    # Step 1 - extract Pandoc's built-in reference.docx (contains all the
    # style names that Pandoc maps its output to).
    pandoc_bin = pypandoc.get_pandoc_path()
    result = subprocess.run(
        [pandoc_bin, "--print-default-data-file", "reference.docx"],
        capture_output=True,
        check=True,
    )
    _REFERENCE_DOCX.write_bytes(result.stdout)

    # Step 2 - open and restyle
    doc = Document(str(_REFERENCE_DOCX))

    heading_color = RGBColor(*_HEADING_COLOR_RGB)
    accent_color = RGBColor(*_ACCENT_RGB)

    # -- Page setup ----------------------------------------------------------
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        _add_page_number_footer(section)

    # -- Restyle existing styles ---------------------------------------------
    heading_sizes = {1: 24, 2: 18, 3: 14, 4: 12, 5: 11, 6: 11}

    for style in doc.styles:
        name = style.name or ""

        # Normal / body text
        if name in ("Normal", "Body Text", "First Paragraph"):
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            pf = style.paragraph_format
            pf.space_after = Pt(6)
            pf.space_before = Pt(0)
            pf.line_spacing = 1.15
            if name == "First Paragraph":
                pf.space_before = Pt(2)

        # Headings 1-6
        elif name.startswith("Heading") and name[-1:].isdigit():
            level = int(name[-1])
            style.font.name = "Calibri"
            style.font.bold = True
            style.font.color.rgb = heading_color
            style.font.size = Pt(heading_sizes.get(level, 11))
            pf = style.paragraph_format
            pf.space_before = Pt(18 if level <= 2 else 12)
            pf.space_after = Pt(6)
            pf.keep_with_next = True
            if level >= 4:
                style.font.bold = False
                style.font.italic = True

        # Source Code (code blocks)
        elif name == "Source Code":
            style.font.name = "Consolas"
            style.font.size = Pt(9.5)
            pf = style.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            pf.line_spacing = 1.0
            _set_paragraph_shading(pf, "F8FAFC")

        # Verbatim Char (inline code)
        elif name == "Verbatim Char":
            style.font.name = "Consolas"
            style.font.size = Pt(9.5)

        # Block Text (block quotes)
        elif name == "Block Text":
            style.font.italic = True
            style.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            pf = style.paragraph_format
            pf.left_indent = Inches(0.4)
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)

        # Hyperlink
        elif name == "Hyperlink":
            style.font.color.rgb = accent_color
            style.font.underline = True

        # Compact (tight lists)
        elif name == "Compact":
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            if style.paragraph_format:
                style.paragraph_format.space_after = Pt(2)

        # Title
        elif name == "Title":
            style.font.name = "Calibri"
            style.font.size = Pt(28)
            style.font.bold = True
            style.font.color.rgb = heading_color
            pf = style.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_after = Pt(4)

        # Subtitle
        elif name == "Subtitle":
            style.font.name = "Calibri"
            style.font.size = Pt(14)
            style.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            pf = style.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_after = Pt(12)

        # Date
        elif name == "Date":
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            style.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            pf = style.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(_REFERENCE_DOCX))


# ---------------------------------------------------------------------------
# OOXML helpers
# ---------------------------------------------------------------------------


def _set_paragraph_shading(paragraph_format, hex_color: str) -> None:
    """Apply background shading to a paragraph style via raw OOXML."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    ppr = paragraph_format._element if hasattr(paragraph_format, "_element") else None
    if ppr is None:
        return
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    ppr.append(shd)


def _add_page_number_footer(section) -> None:
    """Add a centered page number to the section footer via OOXML."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = 1  # CENTER

    run = p.add_run()
    rpr = OxmlElement("w:rPr")
    rsz = OxmlElement("w:sz")
    rsz.set(qn("w:val"), "18")  # 9pt
    rpr.append(rsz)
    rcolor = OxmlElement("w:color")
    rcolor.set(qn("w:val"), "64748B")
    rpr.append(rcolor)
    run._element.append(rpr)

    run.add_text("Page ")

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._element.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char_end)

    run.add_text(" of ")

    fld_char_begin2 = OxmlElement("w:fldChar")
    fld_char_begin2.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin2)

    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    run._element.append(instr2)

    fld_char_end2 = OxmlElement("w:fldChar")
    fld_char_end2.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char_end2)
